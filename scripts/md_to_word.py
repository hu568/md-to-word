#!/usr/bin/env python3
"""
MD 转 Word 转换器 — 基于 Cherry Studio 开源代码修改

Copyright (C) Cherry Studio (https://github.com/CherryHQ/cherry-studio)
Copyright (C) 2026 胡古 (hugo)

本程序是自由软件：您可以根据自由软件基金会发布的
GNU Affero 通用公共许可证（AGPL-3.0）的条款重新分发和/或修改它，
可以选择使用许可证的第3版，或者（按您的选择）任何更高版本。

本程序分发的目的是希望它有用，
但**没有任何担保**，甚至没有隐含的适销性或特定用途的适用性担保。
详情请参阅 GNU Affero 通用公共许可证。

您应该已经随本程序收到了一份 GNU Affero 通用公共许可证的副本。
如果没有，请访问 <https://www.gnu.org/licenses/>。

---

实现说明：
本脚本参考了 Cherry Studio src/main/services/ExportService.ts 的设计思路：
- 使用 markdown-it 系列解析器将 Markdown 解析为 token 流
- 遍历 token 流构建 Word 文档元素
- 表格采用三线表风格设计

修改内容（与原始 TypeScript 版本的区别）：
- 将语言从 TypeScript (Electron) 移植到 Python 3
- 使用 markdown-it-py 替代 markdown-it (JS)
- 使用 python-docx 替代 docx (npm)
- 移除了 Electron 依赖，改为命令行界面
- 移除了 Notion/Obsidian/Joplin 等第三方平台导出功能
- 增加了命令行参数解析和文件读写
"""

import argparse
import os
import re
import sys
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from lxml import etree
from markdown_it import MarkdownIt
from markdown_it.token import Token

# LaTeX 公式转换（md2word 思路：LaTeX → MathML → OMML）
try:
    from latex2mathml.converter import convert as latex_to_mathml
    import mathml2omml
    HAS_LATEX = True
except ImportError:
    HAS_LATEX = False

# 图片后处理模块 — 将 [图片: ...] 占位符替换为真实图片
try:
    from embed_images import process_docx as embed_images_process
    HAS_EMBED_IMAGES = True
except ImportError:
    HAS_EMBED_IMAGES = False

# ── LaTeX 公式支持 ─────────────────────────────────────────────
# 参考 md2word (chenningling/MD2Word) 的转换管线：
#   LaTeX → MathJax/MathML → OMML (Office Math Markup Language)
# 对应 Python 实现：
#   latex2mathml → mathml2omml → python-docx OMML 注入

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
MATH_BLOCK_PREFIX = '§MATH_B'
MATH_INLINE_PREFIX = '§MATH_I'
# 注意：使用 § (U+00A7) 作为前缀，markdown-it-py 视为普通文本
# 注册 OMML 命名空间
etree.register_namespace('m', MATH_NS)


def extract_latex(text: str) -> tuple:
    """
    从 Markdown 中提取 LaTeX 公式，替换为占位符。
    先保护代码块/行内代码，再提取 $$...$$ 和 $...$。

    返回: (处理后的文本, {占位符: {'latex': str, 'display': bool}})
    """
    formulas = {}
    block_n = [0]
    inline_n = [0]

    # 第〇步：保护代码块和行内代码，避免 $$ 和 $ 被误提取
    code_spans = {}  # {占位符: 原始代码}
    cn = [0]

    # 保护代码块 ```...```
    def protect_fence(m):
        key = f'§CODE_{cn[0]}§'
        code_spans[key] = m.group(0)
        cn[0] += 1
        return key

    text = re.sub(r'```.+?```', protect_fence, text, flags=re.DOTALL)

    # 保护行内代码 `...`
    def protect_code(m):
        key = f'§CODE_{cn[0]}§'
        code_spans[key] = m.group(0)
        cn[0] += 1
        return key

    text = re.sub(r'`[^`]+`', protect_code, text)

    # 第一步：块级公式 $$...$$
    def replace_block(m):
        key = f'{MATH_BLOCK_PREFIX}{block_n[0]}§'
        formulas[key] = {'latex': m.group(1).strip(), 'display': True}
        block_n[0] += 1
        return key

    text = re.sub(r'\$\$(.+?)\$\$', replace_block, text, flags=re.DOTALL)

    # 第二步：行内公式 $...$（避开 $$）
    def replace_inline(m):
        key = f'{MATH_INLINE_PREFIX}{inline_n[0]}§'
        formulas[key] = {'latex': m.group(1).strip(), 'display': False}
        inline_n[0] += 1
        return key

    text = re.sub(r'(?<!\$)\$(.+?)\$(?!\$)', replace_inline, text)

    # 第三步：恢复被保护的代码块
    for key, original in code_spans.items():
        text = text.replace(key, original)

    return text, formulas


def latex_to_omml_elem(latex_str: str) -> OxmlElement:
    """将 LaTeX 字符串转换为 OMML OxmlElement"""
    mathml = latex_to_mathml(latex_str)
    omml_str = mathml2omml.convert(mathml)
    if 'xmlns:m' not in omml_str:
        omml_str = omml_str.replace('<m:oMath',
                                     f'<m:oMath xmlns:m="{MATH_NS}"')
    return parse_xml(omml_str)


def insert_inline_equation(paragraph, latex_str: str) -> bool:
    """在段落末尾插入行内公式，返回是否成功"""
    try:
        omml_elem = latex_to_omml_elem(latex_str)
        paragraph._element.append(omml_elem)
        return True
    except Exception as e:
        run = paragraph.add_run(f'[{latex_str}]')
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.size = Pt(10)
        return False


def insert_display_equation(paragraph, latex_str: str) -> bool:
    """插入块级公式（居中 oMathPara），返回是否成功"""
    try:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        omath_para = etree.SubElement(paragraph._element,
                                       f'{{{MATH_NS}}}oMathPara')
        omml_elem = latex_to_omml_elem(latex_str)
        omath_para.append(omml_elem)
        return True
    except Exception as e:
        run = paragraph.add_run(f'[{latex_str}]')
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.size = Pt(10)
        return False


def split_text_by_placeholders(text: str, math_formulas: Dict) -> list:
    """
    将含有数学占位符的文本拆分成段。
    返回: [(文本片段, 是否为占位符)]
    """
    segments = []
    pattern = re.compile(r'§(MATH_[BI]\d+)§')
    last_end = 0
    for m in pattern.finditer(text):
        start, end = m.start(), m.end()
        if start > last_end:
            segments.append((text[last_end:start], False))
        key = m.group(0)
        if key in math_formulas:
            segments.append((math_formulas[key]['latex'], True))
        else:
            segments.append((key, False))
        last_end = end
    if last_end < len(text):
        segments.append((text[last_end:], False))
    return segments


def process_inline_tokens(tokens: List[Token], paragraph, base_font_size: int = 12,
                          base_font_name: str = '微软雅黑', is_header: bool = False,
                          md_source_path: str = None,
                          math_formulas: Dict = None) -> None:
    """
    处理内联 token（粗体、斜体、代码、链接、删除线等）
    参考 Cherry Studio ExportService.processInlineTokens()
    """
    bold_count = 0
    italic_count = 0
    strikethrough = False
    inside_link = False
    link_text = ''
    link_url = ''
    link_start_run_count = 0  # 用于 hyperlink 包装：记录 link_open 时的 run 数量

    i = 0
    while i < len(tokens):
        token = tokens[i]

        if token.type == 'strong_open':
            bold_count += 1
            i += 1
            continue
        elif token.type == 'strong_close':
            bold_count = max(0, bold_count - 1)
            i += 1
            continue
        elif token.type == 'em_open':
            italic_count += 1
            i += 1
            continue
        elif token.type == 'em_close':
            italic_count = max(0, italic_count - 1)
            i += 1
            continue
        elif token.type == 's_open':
            strikethrough = True
            i += 1
            continue
        elif token.type == 's_close':
            strikethrough = False
            i += 1
            continue
        elif token.type == 'link_open':
            inside_link = True
            link_url = ''
            if token.attrs:
                for k, v in token.attrs.items():
                    if k == 'href':
                        link_url = v
            # 记录当前 run 数量，用于 link_close 时包装为真正的超链接
            link_start_run_count = len(paragraph._element.findall(qn('w:r')))
            i += 1
            continue
        elif token.type == 'link_close':
            if inside_link and link_url:
                # 将链接期间添加的 run 包装成 Word 真正的可点击超链接
                all_runs = paragraph._element.findall(qn('w:r'))
                new_runs = all_runs[link_start_run_count:]
                if new_runs:
                    # 添加超链接关系到文档部件
                    r_id = paragraph.part.relate_to(
                        link_url,
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                        is_external=True,
                    )
                    # 获取第一个链接 run 在段落子元素中的位置
                    first_run = new_runs[0]
                    first_run_idx = list(paragraph._element).index(first_run)
                    # 创建超链接 XML 元素
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)
                    # 将 run 从段落移到超链接元素中
                    for run in new_runs:
                        paragraph._element.remove(run)
                        hyperlink.append(run)
                    # 在原来第一个 run 的位置插入超链接元素
                    paragraph._element.insert(first_run_idx, hyperlink)
            inside_link = False
            link_url = ''
            i += 1
            continue
        elif token.type == 'text':
            text = token.content
            # 检查是否包含数学公式占位符
            if math_formulas and re.search(r'§(MATH_[BI]\d+)§', text):
                segments = split_text_by_placeholders(text, math_formulas)
                for seg_text, is_formula in segments:
                    if is_formula:
                        insert_inline_equation(paragraph, seg_text)
                    else:
                        run = paragraph.add_run(seg_text)
                        run.font.size = Pt(base_font_size)
                        run.font.name = base_font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font_name)
                        if bold_count > 0 or is_header:
                            run.bold = True
                        if italic_count > 0:
                            run.italic = True
                        if strikethrough:
                            run.font.strike = True
                        if inside_link and link_url:
                            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                            run.underline = True
            else:
                font_size = Pt(10) if (bold_count > 0 or italic_count > 0) else Pt(base_font_size)
                run = paragraph.add_run(text)
                run.font.size = font_size
                run.font.name = base_font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font_name)
                if bold_count > 0 or is_header:
                    run.bold = True
                if italic_count > 0:
                    run.italic = True
                if strikethrough:
                    run.font.strike = True
                if inside_link and link_url:
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                    run.underline = True
            i += 1
            continue
        elif token.type == 'code_inline':
            run = paragraph.add_run(token.content)
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
            if bold_count > 0 or is_header:
                run.bold = True
            if italic_count > 0:
                run.italic = True
            i += 1
            continue
        elif token.type == 'image':
            # 处理图片
            src = ''
            alt = token.content or ''
            for attr in token.attrs.items() if hasattr(token.attrs, 'items') else []:
                if attr[0] == 'src':
                    src = attr[1]
            if src and os.path.exists(src):
                try:
                    paragraph.add_run().add_picture(src, width=Inches(4))
                except Exception:
                    run = paragraph.add_run(f'[图片: {alt or src}]')
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    run.font.size = Pt(10)
            else:
                # 图片不存在，先尝试相对于MD源文件路径解析
                found = False
                if src and md_source_path:
                    md_dir = os.path.dirname(os.path.abspath(md_source_path))
                    alt_paths = [
                        os.path.join(md_dir, src),
                        os.path.join(md_dir, 'images', os.path.basename(src)),
                    ]
                    for ap in alt_paths:
                        if os.path.exists(ap):
                            try:
                                paragraph.add_run().add_picture(ap, width=Inches(4))
                                found = True
                                break
                            except Exception:
                                pass
                if not found:
                    run = paragraph.add_run(f'[图片: {alt or src}]')
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    run.font.size = Pt(10)
            i += 1
            continue
        elif token.type == 'softbreak' or token.type == 'hardbreak':
            paragraph.add_run('\n')
            i += 1
            continue
        else:
            # 其他未知类型，尝试获取内容
            if token.content:
                run = paragraph.add_run(token.content)
                run.font.size = Pt(base_font_size)
                run.font.name = base_font_name
            i += 1
            continue


def apply_code_block_shading(paragraph):
    """为代码块段落添加灰色背景和边框"""
    pPr = paragraph._element.get_or_add_pPr()
    # 灰色背景
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
    )
    pPr.append(shd)


def convert_markdown_to_docx(markdown_text: str, output_path: str,
                              md_source_path: str = None) -> str:
    """
    核心转换函数
    将 Markdown 文本转换为 Word 文档并保存

    参数:
        markdown_text: Markdown 文本内容
        output_path: 输出 Word 文件路径
        md_source_path: 源 Markdown 文件路径（用于解析图片路径，可选）

    返回: 输出文件路径
    """
    # 预提取 LaTeX 公式（在 Markdown 解析之前）
    math_formulas = {}
    if HAS_LATEX:
        markdown_text, math_formulas = extract_latex(markdown_text)

    # 初始化解析器
    md = MarkdownIt('default', {'maxNesting': 20})
    tokens = md.parse(markdown_text)

    # 创建 Word 文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    i = 0
    while i < len(tokens):
        token = tokens[i]

        # === 标题 ===
        if token.type == 'heading_open':
            level = int(token.tag[1])  # 1-6
            inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            heading_text = inline_tok.content if inline_tok and inline_tok.type == 'inline' else ''
            heading = doc.add_heading('', level=level)

            # 判断标题中是否含公式
            if inline_tok and inline_tok.type == 'inline' and inline_tok.children:
                process_inline_tokens(inline_tok.children, heading,
                                       base_font_size=16 - level,
                                       is_header=True,
                                       math_formulas=math_formulas)
            else:
                run = heading.add_run(heading_text)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            # 设置微软雅黑
            for run in heading.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 3  # 跳过 inline + heading_close
            continue

        # === 段落 ===
        if token.type == 'paragraph_open':
            inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            inline_tokens = inline_tok.children if inline_tok and inline_tok.type == 'inline' and inline_tok.children else []

            # 检查是否为块级公式段落
            if math_formulas and inline_tok and inline_tok.type == 'inline':
                raw_text = inline_tok.content
                block_math_match = re.search(r'§MATH_B\d+§', raw_text)
                if block_math_match:
                    # 纯块级公式段落
                    para = doc.add_paragraph()
                    has_equation = False
                    for seg_text, is_formula in split_text_by_placeholders(raw_text, math_formulas):
                        if is_formula:
                            insert_display_equation(para, seg_text)
                            has_equation = True
                        elif seg_text.strip():
                            run = para.add_run(seg_text)
                            run.font.size = Pt(12)
                            run.font.name = '微软雅黑'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    if has_equation:
                        i += 3
                        continue

            # 图片段落特殊处理
            if inline_tokens and inline_tokens[0].type == 'image':
                # 图片段落特殊处理：支持 ![alt](src)<br>*图注* 格式
                break_pos = None
                for idx, t in enumerate(inline_tokens):
                    if t.type in ('hardbreak', 'softbreak') or (t.type == 'text' and t.content.strip() in ('<br>', '<br/>', '<br />')):
                        break_pos = idx
                        break

                if break_pos is not None and break_pos + 1 < len(inline_tokens):
                    img_tokens = inline_tokens[:break_pos]
                    cap_tokens = inline_tokens[break_pos + 1:]

                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    process_inline_tokens(img_tokens, img_para, md_source_path=md_source_path,
                                           math_formulas=math_formulas)

                    cap_para = doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    process_inline_tokens(cap_tokens, cap_para, base_font_size=10,
                                           md_source_path=md_source_path,
                                           math_formulas=math_formulas)
                    for run in cap_para.runs:
                        run.font.size = Pt(10)
                        if not run.italic:
                            run.italic = True
                else:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    process_inline_tokens(inline_tokens, para, md_source_path=md_source_path,
                                           math_formulas=math_formulas)
            else:
                para = doc.add_paragraph()
                if inline_tokens:
                    process_inline_tokens(inline_tokens, para, md_source_path=md_source_path,
                                           math_formulas=math_formulas)
            i += 3  # 跳过 inline + paragraph_close
            continue

        # === 代码块 ===
        if token.type == 'fence':
            code_text = token.content
            # 创建代码块段落（带背景色和边框）
            for line in code_text.split('\n'):
                code_para = doc.add_paragraph()
                run = code_para.add_run(line if line else ' ')
                run.font.name = 'Consolas'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10)
                apply_code_block_shading(code_para)
            i += 1
            continue

        # === 水平分割线 ===
        if token.type == 'hr':
            para = doc.add_paragraph()
            run = para.add_run('─' * 50)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # === 引用块 ===
        if token.type == 'blockquote_open':
            i += 1
            while i < len(tokens) and tokens[i].type != 'blockquote_close':
                inner = tokens[i]
                if inner.type == 'paragraph_open' and i + 1 < len(tokens):
                    inline_tok = tokens[i + 1]
                    if inline_tok.type == 'inline' and inline_tok.children:
                        para = doc.add_paragraph()
                        # 左边距 + 左边框
                        para.paragraph_format.left_indent = Inches(0.5)
                        process_inline_tokens(inline_tok.children, para,
                                               math_formulas=math_formulas)
                        for run in para.runs:
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                            run.italic = True
                    i += 3
                else:
                    i += 1
            i += 1  # 跳过 blockquote_close
            continue

        # === 无序列表 ===
        if token.type == 'bullet_list_open':
            i += 1
            list_counter = 0
            while i < len(tokens) and tokens[i].type != 'bullet_list_close':
                if tokens[i].type == 'list_item_open':
                    i += 1
                    while i < len(tokens) and tokens[i].type not in ('list_item_close', 'bullet_list_close'):
                        if tokens[i].type == 'paragraph_open' and i + 1 < len(tokens):
                            inline_tok = tokens[i + 1]
                            if inline_tok.type == 'inline' and inline_tok.children:
                                para = doc.add_paragraph(style='List Bullet')
                                process_inline_tokens(inline_tok.children, para, base_font_size=11,
                                                       math_formulas=math_formulas)
                            i += 3
                        elif tokens[i].type == 'inline' and tokens[i].children:
                            para = doc.add_paragraph(style='List Bullet')
                            process_inline_tokens(tokens[i].children, para, base_font_size=11,
                                                   math_formulas=math_formulas)
                            i += 1
                            while i < len(tokens) and tokens[i].type not in ('list_item_close', 'bullet_list_close'):
                                if tokens[i].type == 'paragraph_close':
                                    i += 1
                                    break
                                i += 1
                        else:
                            i += 1
                    if i < len(tokens) and tokens[i].type == 'list_item_close':
                        i += 1
                else:
                    i += 1
            i += 1  # 跳过 bullet_list_close
            continue

        # === 有序列表 ===
        if token.type == 'ordered_list_open':
            i += 1
            while i < len(tokens) and tokens[i].type != 'ordered_list_close':
                if tokens[i].type == 'list_item_open':
                    i += 1
                    while i < len(tokens) and tokens[i].type not in ('list_item_close', 'ordered_list_close'):
                        if tokens[i].type == 'paragraph_open' and i + 1 < len(tokens):
                            inline_tok = tokens[i + 1]
                            if inline_tok.type == 'inline' and inline_tok.children:
                                para = doc.add_paragraph(style='List Number')
                                process_inline_tokens(inline_tok.children, para, base_font_size=11,
                                                       math_formulas=math_formulas)
                            i += 3
                        elif tokens[i].type == 'inline' and tokens[i].children:
                            para = doc.add_paragraph(style='List Number')
                            process_inline_tokens(tokens[i].children, para, base_font_size=11,
                                                   math_formulas=math_formulas)
                            i += 1
                            while i < len(tokens) and tokens[i].type not in ('list_item_close', 'ordered_list_close'):
                                if tokens[i].type == 'paragraph_close':
                                    i += 1
                                    break
                                i += 1
                        else:
                            i += 1
                    if i < len(tokens) and tokens[i].type == 'list_item_close':
                        i += 1
                else:
                    i += 1
            i += 1  # 跳过 ordered_list_close
            continue

        # === 表格 ===
        if token.type == 'table_open':
            table_data = []
            is_header = False
            i += 1
            while i < len(tokens) and tokens[i].type != 'table_close':
                t = tokens[i]
                if t.type == 'thead_open':
                    is_header = True
                    i += 1
                    continue
                elif t.type == 'thead_close':
                    i += 1
                    continue
                elif t.type == 'tbody_open':
                    is_header = False
                    i += 1
                    continue
                elif t.type == 'tbody_close':
                    i += 1
                    continue
                elif t.type == 'tr_open':
                    row_cells = []
                    i += 1
                    while i < len(tokens) and tokens[i].type != 'tr_close':
                        if tokens[i].type in ('th_open', 'td_open'):
                            inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
                            if inline_tok and inline_tok.type == 'inline':
                                row_cells.append({
                                    'text': inline_tok.content,
                                    'children': inline_tok.children,
                                })
                            else:
                                row_cells.append({'text': '', 'children': None})
                            i += 3
                        else:
                            i += 1
                    table_data.append({'cells': row_cells, 'is_header': is_header})
                    if i < len(tokens) and tokens[i].type == 'tr_close':
                        i += 1
                else:
                    i += 1
            i += 1  # 跳过 table_close

            # 创建 Word 表格（三线表风格）
            if table_data:
                num_cols = max((len(row['cells']) for row in table_data), default=0)
                num_rows = len(table_data)
                if num_cols > 0 and num_rows > 0:
                    table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data['cells']):
                            if col_idx < num_cols:
                                cell = table.cell(row_idx, col_idx)
                                cell.text = ''
                                para = cell.paragraphs[0]
                                # 如果有子token（含加粗/斜体等格式），用 process_inline_tokens 处理
                                children = cell_data.get('children') if isinstance(cell_data, dict) else None
                                if children:
                                    process_inline_tokens(
                                        children, para,
                                        base_font_size=11,
                                        base_font_name='微软雅黑',
                                        is_header=row_data['is_header'],
                                        math_formulas=math_formulas,
                                    )
                                else:
                                    cell_text = cell_data.get('text', str(cell_data)) if isinstance(cell_data, dict) else str(cell_data)
                                    run = para.add_run(cell_text)
                                    run.font.size = Pt(11)
                                    run.font.name = '微软雅黑'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                                    if row_data['is_header']:
                                        run.bold = True
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 三线表样式：顶部和底部加粗边框
                    tbl = table._tbl
                    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            continue

        # === 任务列表项（- [ ] 和 - [x]）===
        # markdown-it-py 默认将任务列表识别为普通 bullet_list
        # 因此我们在这个位置不去专门处理，交给 bullet_list_open 处理即可

        i += 1

    # 保存文档
    doc.save(output_path)
    print(f'[OK] Word document saved to: {output_path}')

    # ===== 图片后处理 =====
    # 检查原始 Markdown 是否包含图片引用
    if md_source_path and HAS_EMBED_IMAGES and '[图片:' in markdown_text or '![image' in markdown_text.lower() or '.jpg' in markdown_text.lower() or '.png' in markdown_text.lower():
        has_images = False
        md_lower = markdown_text.lower()
        for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
            if ext in md_lower and '[' in md_lower[:md_lower.index(ext) + 10] if ext in md_lower else False:
                has_images = True
                break
        # 更简单的检查: 是否包含 ![ 语法或 [图片: 占位符
        if '![' in markdown_text:
            has_images = True
        if '[图片:' in markdown_text:
            has_images = True

        if has_images:
            print(f'  [INFO] 检测到图片引用，启动图片后处理...')
            try:
                img_dir = os.path.join(os.path.dirname(os.path.abspath(md_source_path)), 'images')
                if os.path.isdir(img_dir):
                    embed_images_process(output_path, img_dir=img_dir, md_path=md_source_path)
                else:
                    embed_images_process(output_path, md_path=md_source_path)
            except Exception as e:
                print(f'  [WARN] 图片后处理失败: {e}')
    # ====================

    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='将 Markdown 文件转换为 Word 文档 (.docx)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python md_to_word.py note.md
  python md_to_word.py note.md output.docx
  python md_to_word.py --inline "# Hello\\n\\n**bold** text" hello.docx
        """
    )
    parser.add_argument('input', nargs='?', help='输入的 Markdown 文件路径')
    parser.add_argument('output', nargs='?', help='输出的 Word 文件路径（可选，默认在输入文件同目录下）')
    parser.add_argument('--inline', '-i', help='直接传入 Markdown 字符串而不是文件路径')

    args = parser.parse_args()

    if args.inline:
        markdown_text = args.inline.replace('\\n', '\n')
        output_path = args.input or 'output.docx'
        convert_markdown_to_docx(markdown_text, output_path)
    elif args.input and os.path.exists(args.input):
        with open(args.input, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        if args.output:
            output_path = args.output
        else:
            base = os.path.splitext(args.input)[0]
            output_path = base + '.docx'
        convert_markdown_to_docx(markdown_text, output_path, md_source_path=args.input)


if __name__ == '__main__':
    main()
