#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
图片嵌入后处理工具 — 将 DOCX 中的 [图片: ...] 占位符替换为真实图片

用法:
    python embed_images.py <input.docx> [--img-dir <图片目录>] [--md-file <源MD路径>]

流程:
    1. 打开已生成的 DOCX 文件
    2. 扫描所有段落，查找 [图片: xxx] 占位符
    3. 在指定图片目录中查找对应图片文件
    4. 替换占位符为真实嵌入图片 + 图注
"""

import argparse
import os
import re
import sys

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


def extract_image_refs_from_md(md_path):
    """
    从 Markdown 源文件中提取所有图片引用。
    支持的格式:
        ![alt](path)<br>*▲ caption*
        ![alt](path)
        ![alt](path)
        *caption*
    """
    refs = []
    md_dir = os.path.dirname(os.path.abspath(md_path)) if md_path else '.'

    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        print(f'  [WARN] 无法读取MD文件 {md_path}: {e}')
        return refs

    # 模式1: ![alt](path)<br>*▲ caption*
    pattern1 = r'!\[(.+?)\]\((.+?)\)<br>\*\s*▲\s*(.+?)\s*\*'
    for m in re.finditer(pattern1, content):
        alt, path, caption = m.groups()
        refs.append({
            'alt': alt.strip(),
            'path': os.path.join(md_dir, path.strip()),
            'caption': f'\u25b2 {caption.strip()}',
        })

    # 模式2: ![alt](path) 紧接下一行 *caption*
    pattern2 = r'!\[(.+?)\]\((.+?)\)\s*\n\s*\*\s*(.+?)\s*\*'
    for m in re.finditer(pattern2, content):
        alt, path, caption = m.groups()
        # 去重 (避免模式1已匹配)
        if not any(r['alt'] == alt.strip() for r in refs):
            refs.append({
                'alt': alt.strip(),
                'path': os.path.join(md_dir, path.strip()),
                'caption': caption.strip(),
            })

    # 模式3: 只有 ![alt](path) 没有显式图注
    pattern3 = r'!\[(.+?)\]\(((?:images/)?[^)\s]+)\)'
    for m in re.finditer(pattern3, content):
        alt, path = m.groups()
        if not any(r['alt'] == alt.strip() for r in refs):
            refs.append({
                'alt': alt.strip(),
                'path': os.path.join(md_dir, path.strip()),
                'caption': alt.strip(),
            })

    return refs


def find_image_file(img_path, img_dir=None):
    """查找图片文件，支持多种路径方式"""
    # 如果路径直接存在
    if os.path.isfile(img_path):
        return img_path

    # 如果指定了图片目录，在目录中找文件名
    if img_dir:
        fname = os.path.basename(img_path)
        candidate = os.path.join(img_dir, fname)
        if os.path.isfile(candidate):
            return candidate

    # 在输入文件同级的 images/ 目录中找
    base_dir = os.path.dirname(os.path.abspath(img_path))
    for subdir in ['', 'images', 'img', 'assets']:
        fname = os.path.basename(img_path)
        candidate = os.path.join(base_dir, subdir, fname)
        if os.path.isfile(candidate):
            return candidate

    return None


def insert_image_at_paragraph(doc, para, img_path, caption):
    """在指定段落处插入图片和图注，返回是否成功"""
    if not os.path.isfile(img_path):
        print(f'    [跳过] 图片不存在: {img_path}')
        return False

    # 清空原段落
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 插入图片 (宽度5英寸)
    run = para.add_run()
    run.add_picture(img_path, width=Inches(5.0))

    # 创建图注段落（紧跟图片段落之后）
    caption_para_elem = OxmlElement('w:p')
    para._element.addnext(caption_para_elem)

    # 构建图注内容
    from docx import Document as TmpDoc
    tmp_doc = TmpDoc()
    tmp_p = tmp_doc.add_paragraph()
    tmp_r = tmp_p.add_run(caption)
    tmp_r.font.size = Pt(9)
    tmp_r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    tmp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for child in tmp_p._element:
        caption_para_elem.append(child)

    return True


def process_docx(docx_path, img_dir=None, md_path=None, dry_run=False):
    """
    处理 DOCX 文件，替换图片占位符。

    参数:
        docx_path: Word 文件路径
        img_dir: 图片目录（可选）
        md_path: 源 Markdown 文件路径（可选，用于提取图注）
        dry_run: 仅扫描不修改

    返回:
        (处理数量, 总占位符数量)
    """
    doc = Document(docx_path)

    # 从MD中提取图片引用
    md_refs = []
    if md_path:
        md_refs = extract_image_refs_from_md(md_path)
        print(f'  从MD中发现 {len(md_refs)} 个图片引用')

    # 扫描所有段落
    placeholders = []
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # 匹配 [图片: xxx] 或 [图片: xxx.jpg] 格式
        match = re.search(r'\[图片:\s*(.+?)\]', text)
        if match:
            placeholders.append((p_idx, para, match.group(1).strip()))

    if not placeholders:
        print('  [跳过] 未发现图片占位符')
        return 0, 0

    print(f'  发现 {len(placeholders)} 个图片占位符')

    if dry_run:
        for p_idx, _, ref in placeholders:
            print(f'    段落[{p_idx}]: {ref}')
        return 0, len(placeholders)

    # 处理每个占位符
    modified = 0
    for p_idx, para, ref in placeholders:
        print(f'  处理段落[{p_idx}]: {ref}', end='')

        # 尝试从MD引用中匹配
        matched = False
        for md_ref in md_refs:
            # 匹配条件: alt文本包含引用文本，或反之
            if (md_ref['alt'] in ref or ref in md_ref['alt'] or
                os.path.basename(md_ref['path']) in ref or
                ref in os.path.basename(md_ref['path'])):
                img_path = find_image_file(md_ref['path'], img_dir)
                if img_path and insert_image_at_paragraph(doc, para, img_path, md_ref['caption']):
                    print(f' -> {os.path.basename(img_path)}')
                    modified += 1
                    matched = True
                    break

        if not matched:
            # 直接尝试把ref当文件名查找
            for try_path in [ref, ref + '.jpg', ref + '.png', ref + '.jpeg']:
                img_path = find_image_file(try_path, img_dir)
                if img_path and insert_image_at_paragraph(doc, para, img_path, f'\u25b2 {ref}'):
                    print(f' -> {os.path.basename(img_path)} (自动匹配)')
                    modified += 1
                    matched = True
                    break

        if not matched:
            print(' [未找到图片]')

    if modified > 0:
        doc.save(docx_path)
        print(f'\n  [OK] 已保存: {docx_path}')

    return modified, len(placeholders)


def main():
    parser = argparse.ArgumentParser(
        description='将 DOCX 中的 [图片: ...] 占位符替换为真实嵌入图片',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 基本用法（自动从MD提取图注信息）
  python embed_images.py output.docx --md-file source.md

  # 指定图片目录
  python embed_images.py output.docx --img-dir ./images --md-file source.md

  # 仅扫描，不修改
  python embed_images.py output.docx --dry-run

  # 从 md_to_word.py 集成调用
  python -c "from embed_images import process_docx; process_docx('out.docx', img_dir='images', md_path='source.md')"
        """
    )
    parser.add_argument('input', help='输入的 Word 文件路径 (.docx)')
    parser.add_argument('--img-dir', help='图片所在目录（默认自动查找）')
    parser.add_argument('--md-file', help='源 Markdown 文件路径（用于提取图注）')
    parser.add_argument('--dry-run', action='store_true', help='仅扫描，不修改文档')
    parser.add_argument('--quiet', '-q', action='store_true', help='静默模式')

    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f'[ERROR] 文件不存在: {args.input}')
        sys.exit(1)

    if not args.quiet:
        print(f'处理文档: {args.input}')

    processed, total = process_docx(
        args.input,
        img_dir=args.img_dir,
        md_path=args.md_file,
        dry_run=args.dry_run,
    )

    if not args.quiet:
        print(f'\n[OK] 完成! 处理 {processed}/{total} 张图片')


if __name__ == '__main__':
    main()
